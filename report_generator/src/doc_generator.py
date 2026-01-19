import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportToDocxConverter:
    """
    Converts a Markdown report to a formatted DOCX file.
    Handles basic Markdown syntax: headers (#), bold (**), italics (*), lists (-), and images (![]()).
    """

    def __init__(self, md_path: Path, output_dir: Path):
        self.md_path = md_path
        self.output_dir = output_dir
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles for a professional look."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Title Style (H1)
        h1_style = self.doc.styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Calibri Light'
        h1_font.size = Pt(24)
        h1_font.color.rgb = RGBColor(44, 62, 80)  # Dark Blue

        # Heading 2 Style
        h2_style = self.doc.styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Calibri Light'
        h2_font.size = Pt(18)
        h2_font.color.rgb = RGBColor(52, 152, 219) # Blue
        
        # Heading 3 Style
        h3_style = self.doc.styles['Heading 3']
        h3_font = h3_style.font
        h3_font.name = 'Calibri'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(127, 140, 141) # Grey

    def convert(self) -> Path:
        """Parse the markdown file and generate DOCX."""
        if not self.md_path.exists():
            raise FileNotFoundError(f"Markdown report not found at {self.md_path}")

        unique_id = self.md_path.stem.replace('event_report_', '')
        output_filename = f"event_report_{unique_id}.docx"
        output_path = self.output_dir / output_filename

        with open(self.md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Handle Headers
            if line.startswith('# '):
                self.doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], level=3)
            
            # Handle Images
            elif line.startswith('![') and '](' in line:
                # Extract image path: ![Alt](path)
                match = re.search(r'\!\[.*\]\((.*)\)', line)
                if match:
                    img_filename = match.group(1)
                    img_path = self.output_dir / img_filename
                    if img_path.exists():
                        try:
                            self.doc.add_picture(str(img_path), width=Inches(6))
                            last_paragraph = self.doc.paragraphs[-1] 
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            self.doc.add_paragraph(f"[Image: {img_filename} - Error inserting]")
            
            # Handle Horizontal Rules
            elif line.startswith('---'):
                self.doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Handle Lists
            elif line.startswith('- ') or line.startswith('* '):
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, line[2:])
            
            # Handle Numbered Lists (approximate)
            elif re.match(r'^\d+\. ', line):
                p = self.doc.add_paragraph(style='List Number')
                text = re.sub(r'^\d+\. ', '', line)
                self._add_formatted_text(p, text)

            # Handle Normal Text
            else:
                p = self.doc.add_paragraph()
                self._add_formatted_text(p, line)

        self.doc.save(str(output_path))
        return output_path

    def _add_formatted_text(self, paragraph, text):
        """Parses bold (**) and italics (*) within a line."""
        # This is a basic parser. It splits by ** first.
        # A more robust solution would be a proper tokenizer, but this covers 90% of cases.
        
        # Split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Split non-bold parts by italics
                italic_parts = re.split(r'(\*.*?\*)', part)
                for subpart in italic_parts:
                    if subpart.startswith('*') and subpart.endswith('*'):
                        run = paragraph.add_run(subpart[1:-1])
                        run.italic = True
                    else:
                        paragraph.add_run(subpart)
