import docx
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging
import re

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class MissingTemplateAsset(Exception):
    """Raised when the template requires an asset (image) that was not provided."""
    def __init__(self, missing_assets: List[Dict[str, Any]]):
        super().__init__("Missing template assets: " + ", ".join([m['marker'] for m in missing_assets]))
        self.missing_assets = missing_assets


class MissingTemplatePlaceholder(Exception):
    """Raised when the template contains mandatory placeholders that were not found in the provided template."""
    def __init__(self, missing_placeholders: List[str]):
        super().__init__("Missing placeholders: " + ", ".join(missing_placeholders))
        self.missing_placeholders = missing_placeholders


class DocxReportGenerator:
    """
    Generates strict format reports based on a provided .docx template.
    Handles text replacement, image insertion, and table population.

    Key improvements:
    - Preserves font/style/size of the template label when inserting replacement text
    - Parses optional image directives in the template for width/height/alignment (e.g. "Event Poster: [w=15cm;h=10cm]")
    - Raises a `MissingTemplateAsset` when the template requires images that are not provided so the caller can prompt the user for input
    - Validates that required placeholder labels exist in the template and raises `MissingTemplatePlaceholder` if they don't
    """

    IMAGE_DIRECTIVE_RE = re.compile(r"\[(.+?)\]")  # content inside [] e.g. w=15cm;h=10cm;align=center

    def __init__(self, template_path: Path):
        self.template_path = template_path
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found at {self.template_path}")

    def generate_report(self, data: Dict[str, Any], output_path: Path, charts: Dict[str, Path]):
        """
        Main method to generate the report.

        Args:
            data: Dictionary of data to populate (text fields, stats).
            output_path: Path to save the generated report.
            charts: Dictionary mapping chart keys (e.g., 'ratings', 'demographics', 'poster') to file paths.

        Returns:
            True on success. Raises exceptions for missing required template items.
        """
        logger.info(f"Generating report using template: {self.template_path}")
        doc = docx.Document(self.template_path)

        # Map template labels to data keys (same as before)
        label_map = {
            "Title:": data.get('event_name', ''),
            "Name of the Event": data.get('event_name', ''),
            "Date:": data.get('date', '2025-03-15'),
            "Time:": data.get('time', '10:00 AM - 5:00 PM'),
            "Venue:": data.get('venue', 'Main Auditorium'),
            "Target Audience:": data.get('target_audience', 'Students & Faculty'),
            "No. of Participants Present:": str(data.get('total_participants', 0)),
            "No. of Girl Participants Present:": str(data.get('female_count', 0)),
            "No. of Boy Participants Present:": str(data.get('male_count', 0)),
            "Resource Person:": data.get('resource_person', 'Dr. Alan Turing'),
            "Organization of Recourse Person:": data.get('rp_org', 'AI Research Lab'),
            "Organizing Department": data.get('department', 'Computer Engineering'),
            "Faculty Coordinator:": data.get('coordinator', 'Prof. Smith'),
            "Objectives:": self._format_list(data.get('objectives', ['Learn AI', 'Build Networks'])),
            "Outcomes:": self._format_list(data.get('outcomes', ['Understanding of LLMs', 'Practical skills'])),
            "Detailed Report:": data.get('detailed_report', 'The event proceeded with an inauguration followed by technical sessions...'),
            "Department Name": data.get('institution_name', 'Department of Computer Engineering'),
            "Social Media Links:": data.get('social_links', 'Facebook: @techfest, Instagram: @techfest_ai')
        }

        # Ensure placeholders exist in the template if strict mode desired
        missing_placeholders = self._find_missing_placeholders(doc, list(label_map.keys()))
        if missing_placeholders:
            # If placeholders are missing, raise and let caller decide how to proceed
            raise MissingTemplatePlaceholder(missing_placeholders)

        # 1. Text Replacement (preserving runs/styles)
        self._process_text_replacements(doc, label_map)

        # 2. Image Insertion (Charts, Poster, Snapshots). This may raise MissingTemplateAsset.
        self._process_image_insertions(doc, charts)

        # 3. Table Population (Student List)
        if 'student_list_table' in data:
            self._populate_student_table(doc, data['student_list_table'])

        logger.info(f"Saving report to: {output_path}")
        doc.save(output_path)
        return True

    def _find_missing_placeholders(self, doc: docx.Document, labels: List[str]) -> List[str]:
        """Return list of labels that are not found anywhere in the document (paragraphs or table cells)."""
        found = set()
        for p in doc.paragraphs:
            for label in labels:
                if label in p.text:
                    found.add(label)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for label in labels:
                            if label in p.text:
                                found.add(label)
        missing = [l for l in labels if l not in found]
        return missing

    def _copy_run_formatting(self, src_run, dest_run):
        """Copies basic formatting from src_run to dest_run."""
        try:
            dest_run.bold = src_run.bold
            dest_run.italic = src_run.italic
            dest_run.underline = src_run.underline
            dest_run.font.name = src_run.font.name
            dest_run.font.size = src_run.font.size
            dest_run.font.color.rgb = src_run.font.color.rgb
        except Exception:
            # Silently ignore unsupported formatting copy in some cases
            pass

    def _process_text_replacements(self, doc: docx.Document, label_map: Dict[str, Any]):
        """Replace labels preserving the formatting of the label run.

        This function looks for runs that contain a label and injects a new run after it with the
        replacement text and copies the formatting from the label run to preserve size and style.
        """

        def replace_in_paragraph(p):
            # Iterate a copy of the runs since we may mutate runs while iterating
            runs = list(p.runs)
            for i, r in enumerate(runs):
                for label, value in label_map.items():
                    if label in r.text:
                        # Split run text around the label
                        parts = r.text.split(label)
                        prefix = parts[0]
                        suffix = label.join(parts[1:]) if len(parts) > 1 else ''

                        # Keep the current run as prefix
                        r.text = prefix + label

                        # Insert replacement run after the current run
                        insert_index = i + 1
                        new_run = p.add_run()  # append at end then we will reorder if necessary
                        new_run.text = f" {value}"
                        self._copy_run_formatting(r, new_run)

                        # Append remaining suffix as a new run to maintain existing text
                        if suffix:
                            suffix_run = p.add_run(suffix)
                            # suffix should maintain original formatting as well
                            self._copy_run_formatting(r, suffix_run)

                        # We limit replacement to first occurrence within this run
                        break

        for p in doc.paragraphs:
            replace_in_paragraph(p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

    def _format_list(self, items: List[str]) -> str:
        """Formats a list of strings into a newline-separated string with bullets."""
        return "\n" + "\n".join([f"❖ {item}" for item in items])

    def _parse_image_directives(self, text: str) -> Dict[str, Any]:
        """Parses directive in square brackets and returns width/height/align if provided.

        Example: "Event Poster: [w=15cm;h=10cm;align=center]"
        Returns: { 'width_cm': 15.0, 'height_cm': 10.0, 'align': 'center' }
        """
        match = self.IMAGE_DIRECTIVE_RE.search(text)
        if not match:
            return {}
        content = match.group(1)
        props = {}
        for kv in content.split(';'):
            if '=' in kv:
                k, v = kv.split('=', 1)
                k = k.strip().lower()
                v = v.strip()
                if k in ('w', 'width') and v.endswith('cm'):
                    props['width_cm'] = float(v[:-2])
                elif k in ('h', 'height') and v.endswith('cm'):
                    props['height_cm'] = float(v[:-2])
                elif k == 'align':
                    props['align'] = v.lower()
        return props

    def _process_image_insertions(self, doc: docx.Document, charts: Dict[str, Path]):
        """
        Inserts images at specific markers. If the template contains an image marker but the image
        is missing from `charts`, raises `MissingTemplateAsset` with details so the caller can prompt the user.
        """
        markers = {
            "Logo:": charts.get('logo'),
            "Organization Logo:": charts.get('logo'),
            "Snapshot of the Event:": charts.get('snapshot'),
            "Feedback Analysis:": charts.get('ratings_chart'),
            "Event Poster:": charts.get('poster')
        }

        missing_assets = []

        for p in doc.paragraphs:
            for marker, image_path in markers.items():
                if marker in p.text:
                    # Parse any size directives from the same paragraph
                    directives = self._parse_image_directives(p.text)

                    if not image_path or not Path(image_path).exists():
                        # If an image is required by the template (marker exists), but not provided, collect it
                        missing_assets.append({
                            'marker': marker,
                            'directives': directives
                        })
                        continue

                    # Insert image with the required size if provided
                    r = p.add_run()
                    r.add_break()
                    try:
                        # Choose a sensible default logo width if the marker appears to be a logo
                        default_width = 4 if 'logo' in marker.lower() else 15
                        width = Cm(directives.get('width_cm', default_width))
                        # Only set height if explicitly specified
                        if 'height_cm' in directives:
                            r.add_picture(str(image_path), width=width, height=Cm(directives['height_cm']))
                        else:
                            r.add_picture(str(image_path), width=width)

                        # Align paragraph if requested
                        if directives.get('align') == 'center':
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        logger.info(f"Inserted image for {marker} with directives {directives}")
                    except Exception as e:
                        logger.error(f"Failed to insert image {image_path} for marker {marker}: {e}")

        if missing_assets:
            raise MissingTemplateAsset(missing_assets)

    def _populate_student_table(self, doc: docx.Document, student_list: List[Dict[str, Any]]):
        """
        Finds the student table (header: "S. No.", "Name", "Branch") and adds rows.
        """
        target_table = None

        # Find the correct table by header
        for table in doc.tables:
            if len(table.rows) > 0:
                header_row = table.rows[0]
                cells = [c.text.strip() for c in header_row.cells]
                if "S. No." in cells and "Name" in cells:
                    target_table = table
                    break

        if target_table:
            for student in student_list:
                row = target_table.add_row()
                # Assuming simple 3 column structure based on prompt
                # S. No. | Name | Branch
                if len(row.cells) >= 3:
                    row.cells[0].text = str(student['s_no'])
                    row.cells[1].text = str(student['name'])
                    row.cells[2].text = str(student['branch'])
            logger.info(f"Populated student table with {len(student_list)} rows")
        else:
            logger.warning("Student table header not found in template.")

