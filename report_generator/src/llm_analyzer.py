import ollama
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from enum import Enum

class AnalysisType(Enum):
    """Types of event feedback analysis."""
    POSITIVE = "positive"
    NEGATIVE = "negative"
    SOCIAL = "social"

@dataclass
class LLMConfig:
    """Configuration for LLM-based analysis in event management system."""
    model_name: str = "llama3:8b"
    temperature: float = 0.7
    max_retries: int = 2
    timeout: int = 30

class EventFeedbackAnalyzer:
    """
    Handles AI-powered analysis of event feedback and social media sentiment.
    
    This class uses Large Language Models (LLMs) to automatically analyze
    qualitative feedback from college event participants, identifying key
    themes, sentiment, and areas for improvement.
    """
    
    def __init__(self, config: Optional[LLMConfig] = None):
        """
        Initialize the Event Feedback Analyzer.
        
        Args:
            config: Optional LLM configuration settings
        """
        self.config = config or LLMConfig()
        
    def _format_comments(self, comments: List[str]) -> str:
        """Format participant comments into a bulleted list."""
        return "\n- ".join(comments)
    
    def _format_social_posts(self, posts: List[Dict]) -> str:
        """Format social media posts with sentiment labels."""
        return "\n- ".join([
            f"[{post.get('sentiment', 'neutral').upper()}] {post.get('text', '')}"
            for post in posts
        ])
    
    def _build_prompt(self, analysis_type: AnalysisType, content: str, event_details: Optional[Dict] = None) -> str:
        """Build appropriate prompt based on analysis type and event context."""
        event_details = event_details or {}
        # Get event details with fallbacks
        event_name = event_details.get('name', 'a college tech event')
        event_type = event_details.get('type', 'workshop/hackathon')
        
        prompts = {
            AnalysisType.POSITIVE: f"""You are an AI assistant analyzing feedback for a college event.
Event Name: {event_name}
Event Type: {event_type}

Review the feedback and list the top 3-5 POSITIVE themes participants enjoyed.
Focus on: content, speakers, organization, networking, hands-on activities.

Participant Feedback:
- {content}

Key Positive Themes:""",
            
            AnalysisType.NEGATIVE: f"""You are an AI assistant analyzing feedback for a college event.
Event Name: {event_name}
Event Type: {event_type}

Review the feedback and list the top 3-5 NEGATIVE themes or areas for improvement.
Focus on: content gaps, timing issues, technical problems, logistics.

Participant Feedback:
- {content}

Areas for Improvement:""",
            
            AnalysisType.SOCIAL: f"""You are an AI assistant analyzing social media posts for a college event.
Event Name: {event_name}
Event Type: {event_type}

Summarize the overall sentiment in 2-3 sentences, highlighting:
- General sentiment (positive/negative/mixed)
- Most discussed topics or highlights
- Any notable concerns or praises
- Student engagement level

Social Media Posts:
- {content}

Social Media Sentiment Summary:"""
        }
        return prompts[analysis_type]
    
    def _call_llm(self, prompt: str, context: str = "") -> str:
        """
        Call Ollama LLM with retry logic and error handling.
        
        Args:
            prompt: The prompt to send to the LLM
            context: Context description for logging (e.g., "positive feedback")
            
        Returns:
            The LLM response content
        """
        for attempt in range(1, self.config.max_retries + 1):
            try:
                if attempt > 1:
                    print(f"  🔄 Retry attempt {attempt}/{self.config.max_retries}...")
                
                response = ollama.chat(
                    model=self.config.model_name,
                    messages=[{'role': 'user', 'content': prompt}],
                    options={
                        'temperature': self.config.temperature,
                    }
                )
                
                content = response.get('message', {}).get('content', '').strip()
                
                if not content:
                    raise ValueError("Empty response from LLM")
                
                return content
                
            except ollama.ResponseError as e:
                print(f"  ⚠️  Ollama response error on attempt {attempt}: {e}")
                if attempt == self.config.max_retries:
                    return f"Error: Unable to analyze after {self.config.max_retries} attempts."
                    
            except Exception as e:
                print(f"  ❌ ERROR on attempt {attempt}: {e}")
                if attempt == self.config.max_retries:
                    error_msg = "Error: Could not connect to AI analysis service."
                    print(f"  💡 TIP: Ensure Ollama is running (`ollama serve`)")
                    print(f"  💡 TIP: Verify model '{self.config.model_name}' is installed")
                    return error_msg
        
        return "Error: Maximum retry attempts exceeded."
    
    def analyze_event_feedback(
        self, 
        comments: List[str],
        event_details: Optional[Dict] = None
    ) -> Tuple[str, str]:
        """
        Analyze event feedback comments to identify positive and negative themes.
        Accepts optional event_details to provide context (name/type) to the LLM.
        """
        if not comments:
            return "No feedback comments provided.", "No feedback comments provided."
        
        # Filter out empty comments
        valid_comments = [c.strip() for c in comments if c and c.strip()]
        
        if not valid_comments:
            return "No valid feedback to analyze.", "No valid feedback to analyze."
        
        formatted_comments = self._format_comments(valid_comments)
        
        print(f"🤖 AI Analysis: Processing {len(valid_comments)} feedback comments...")
        print(f"  → Identifying positive themes...")
        positive_prompt = self._build_prompt(AnalysisType.POSITIVE, formatted_comments, event_details)
        positive_summary = self._call_llm(positive_prompt, "positive feedback")
        
        print(f"  → Identifying improvement areas...")
        negative_prompt = self._build_prompt(AnalysisType.NEGATIVE, formatted_comments, event_details)
        negative_summary = self._call_llm(negative_prompt, "improvement areas")
        
        print(f"  ✅ Feedback analysis complete")
        
        return positive_summary, negative_summary
    
    def analyze_social_sentiment(
        self, 
        social_posts: List[Dict[str, str]]
    ) -> str:
        """
        Analyze social media posts about the event.
        
        This method uses AI to understand public sentiment and engagement
        around the event on social media platforms.
        
        Args:
            social_posts: List of social media post dictionaries with 'text' and 'sentiment'
            
        Returns:
            Summary of social media sentiment and engagement
        """
        if not social_posts:
            return "No social media data available for this event."
        
        # Filter valid posts
        valid_posts = [
            post for post in social_posts 
            if isinstance(post, dict) and post.get('text', '').strip()
        ]
        
        if not valid_posts:
            return "No valid social media posts to analyze."
        
        formatted_posts = self._format_social_posts(valid_posts)
        
        print(f"🤖 AI Analysis: Processing {len(valid_posts)} social media posts...")
        prompt = self._build_prompt(AnalysisType.SOCIAL, formatted_posts)
        
        summary = self._call_llm(prompt, "social media")
        print(f"  ✅ Social media analysis complete")
        
        return summary
    
    def generate_recommendations(
        self, 
        event_stats: Dict[str, any],
        positive_themes: str,
        improvement_areas: str
    ) -> str:
        """
        Generate AI-powered recommendations for future events.
        
        Args:
            event_stats: Dictionary containing event statistics
            positive_themes: Summary of positive feedback themes
            improvement_areas: Summary of areas needing improvement
            
        Returns:
            AI-generated recommendations for future events
        """
        context = f"""You are an AI assistant helping college event organizers improve future tech events.

Based on the event data below, provide 4-6 specific, actionable recommendations for organizing better events in the future.

Event Statistics:
- Total Participants: {event_stats.get('total_participants', 'N/A')}
- Average Session Rating: {event_stats.get('avg_rating', 'N/A'):.2f}/5
- Student Participation: {event_stats.get('student_count', 'N/A')} students

Key Strengths:
{positive_themes}

Areas for Improvement:
{improvement_areas}

Focus your recommendations on:
1. Enhancing the positive aspects
2. Addressing the identified issues
3. Improving student engagement
4. Better resource utilization
5. Technical/logistical improvements

Provide recommendations in a clear, prioritized bullet-point format.

Recommendations for Future Events:"""
        
        print(f"🤖 AI Analysis: Generating event recommendations...")
        recommendations = self._call_llm(context, "recommendations")
        print(f"  ✅ Recommendations generated")
        
        return recommendations

    def generate_detailed_report(self, hints: str, event_name: str) -> str:
        """Generate a single cohesive paragraph for the detailed report section."""
        if not hints or not hints.strip():
            return f"The {event_name} was conducted successfully with enthusiastic participation."
            
        prompt = f"""You are writing a formal academic report for an event named "{event_name}".
Based on the following pointers provided by the organizer, write exactly ONE cohesive, formal paragraph summarizing the event details.
Do NOT include any greetings, bullet points, or extra commentary. Just the paragraph.

Pointers:
{hints}

Detailed Report Paragraph:"""
        return self._call_llm(prompt, "detailed report")

    def generate_feedback_summary_text(self, comments: List[str], event_name: str) -> str:
        """Generate a concise textual summary of feedback for the template."""
        if not comments:
            return "No feedback comments provided by participants."
            
        valid_comments = [c.strip() for c in comments if c and c.strip()]
        if not valid_comments:
            return "No valid feedback comments provided by participants."
            
        formatted_comments = self._format_comments(valid_comments)
        prompt = f"""You are summarizing feedback for a formal academic event report for "{event_name}".
Based on the following participant comments, write a short, cohesive summary (1-2 paragraphs max) highlighting the general sentiment, key positive takeaways, and any notable constructive criticism.
Do NOT use bullet points or markdown formatting. Keep it formal and academic.

Participant Feedback:
{formatted_comments}

Feedback Summary:"""
        return self._call_llm(prompt, "feedback summary")

    def fill_docx_template(self, template_path: 'Path', context_data: Dict[str, any]) -> Dict[str, str]:
        """
        Instruct the LLM to strictly fill a .docx template.

        The method extracts a textual skeleton of the template and provides it along with
        the structured context_data. The LLM is instructed to return a JSON object only,
        where keys are exact placeholder labels found in the template and values are the
        corresponding text/image filenames to be inserted.
        """
        import json
        from pathlib import Path
        from docx import Document

        template_path = Path(template_path)
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        # Extract a textual skeleton from the docx: paragraphs + table cell texts
        doc = Document(str(template_path))
        parts = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text.strip())

        skeleton = "\n".join(parts)

        prompt = f"""
You are an assistant that must STRICTLY fill a Word (.docx) template.

Below is the textual skeleton of the uploaded template (extracted from the .docx). The skeleton contains labels and any image markers.

Template Skeleton:
{skeleton}

User Data (JSON):
{json.dumps(context_data, ensure_ascii=False, indent=2)}

Instructions (IMPORTANT):
- Return ONLY a valid JSON object and nothing else.
- Keys in the JSON must match the exact labels/placeholders in the template skeleton (case-sensitive) that should be replaced.
- For textual placeholders return plain strings containing the content to insert.
- For image placeholders return a filename (e.g., "logo.png" or "poster.png") corresponding to an uploaded image.
- Do NOT modify the template structure or return the entire document; return only the mapping for placeholders.

Provide the JSON mapping now."""

        # Call LLM with deterministic settings
        orig_temp = self.config.temperature
        self.config.temperature = 0.0
        try:
            response = self._call_llm(prompt, "fill_docx_template")
        finally:
            self.config.temperature = orig_temp

        # Parse JSON
        try:
            # Some LLM responses may be wrapped in Markdown code fences; remove them
            resp = response.strip()
            if resp.startswith('```'):
                # remove triple-backtick wrappers and optional language tag
                parts = resp.split('```')
                if len(parts) >= 3:
                    resp = parts[1]
                else:
                    resp = resp.replace('```', '')
            # Locate the JSON object between the first { and last }
            first = resp.find('{')
            last = resp.rfind('}')
            resp_candidate = resp[first:last+1] if first != -1 and last != -1 else resp

            # Some LLMs place multiline content inside triple-quoted strings ("""...""") which is
            # invalid JSON. Convert any triple-quoted blocks into proper JSON string literals.
            import re
            def _replace_triple_quotes(match):
                inner = match.group(1)
                # Use json.dumps to properly escape newlines/quotes
                return json.dumps(inner)

            resp_candidate = re.sub(r'"""(.*?)"""', _replace_triple_quotes, resp_candidate, flags=re.S)

            parsed = json.loads(resp_candidate)
            if not isinstance(parsed, dict):
                raise ValueError("LLM did not return a JSON object as expected")
            return parsed
        except Exception as e:
            raise ValueError(f"Failed to parse LLM response as JSON: {e}\nResponse content: {response}")

    def fill_tex_template(self, template_path: 'Path', context_data: Dict[str, any]) -> str:
        """
        Instruct the LLM to strictly fill a .tex (LaTeX) template and return the full
        filled .tex content as plain text (the same format as the uploaded template).
        """
        import json
        from pathlib import Path

        template_path = Path(template_path)
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        template_text = template_path.read_text(encoding='utf-8')

        prompt = f"""
You are an assistant that must STRICTLY fill a LaTeX (.tex) template.

Below is the EXACT LaTeX template content. Fill placeholders using the User Data, and return ONLY the complete filled .tex source. Do not include any explanations or extra text.

TEMPLATE:
{template_text}

USER DATA (JSON):
{json.dumps(context_data, ensure_ascii=False, indent=2)}

RETURN the filled LaTeX document now, and nothing else.
"""
        # Deterministic call
        orig_temp = self.config.temperature
        self.config.temperature = 0.0
        try:
            response = self._call_llm(prompt, "fill_tex_template")
        finally:
            self.config.temperature = orig_temp

        return response


# Convenience functions for backward compatibility
def get_llm_summary(comments_list: List[str], model_name: str = "llama3:8b") -> Tuple[str, str]:
    """
    Legacy function for feedback analysis.
    
    Args:
        comments_list: List of feedback comments
        model_name: Ollama model to use
        
    Returns:
        Tuple of (positive_summary, improvement_areas)
    """
    config = LLMConfig(model_name=model_name)
    analyzer = EventFeedbackAnalyzer(config)
    return analyzer.analyze_event_feedback(comments_list)

def get_social_summary(social_list: List[Dict], model_name: str = "llama3:8b") -> str:
    """
    Legacy function for social media analysis.
    
    Args:
        social_list: List of social media posts
        model_name: Ollama model to use
        
    Returns:
        Social media sentiment summary
    """
    config = LLMConfig(model_name=model_name)
    analyzer = EventFeedbackAnalyzer(config)
    return analyzer.analyze_social_sentiment(social_list)


# Example usage and testing
if __name__ == "__main__":
    print("🧪 Testing Event Feedback Analyzer...\n")
    
    # Initialize analyzer
    config = LLMConfig(model_name="llama3:8b", temperature=0.7)
    analyzer = EventFeedbackAnalyzer(config)
    
    # Test feedback analysis
    test_comments = [
        "The AI workshop was fantastic! Great hands-on experience.",
        "Speakers were knowledgeable but sessions ran late.",
        "Loved the networking session with industry professionals.",
        "Food was cold and WiFi kept disconnecting during demos.",
        "Amazing opportunity to learn about machine learning!"
    ]
    
    print("="*60)
    positive, negative = analyzer.analyze_event_feedback(test_comments)
    print("\n📊 POSITIVE THEMES:")
    print(positive)
    print("\n⚠️  IMPROVEMENT AREAS:")
    print(negative)
    print("="*60)